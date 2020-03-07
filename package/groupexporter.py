#   Author : github.com/slingthy
#   Time : March 4, 2020
#   ---------------------------------------------
#   The purpose is to extract someone's record 
#   from wechat group and store it in a docx file.
#   !!!JUST 'TEXT' message.
#   ---------------------------------------------
#   Before RUN:
#   1. COPY the WeChat Group message
#   2. PASTE into Word document as 'xxxxxx.docx'
#   3. SAVE it, do NOT make any change
#   ---------------------------------------------
#   Dependencies : python-docx, lxml

from docx import Document
from time import strftime,localtime
import os

def pymain(pathroad,groupmember,savename):
    when=strftime("%m.%d", localtime())
    D = Document(pathroad)
    D2= Document()
    para=D.paragraphs
    

    #
    for i in range(len(para)):
        if para[i].text.find(groupmember)==-1:
            raise ValueError
    if savename=='':
        raise ValueError
    for i in range(len(para)):
        if para[i].text.find(":")>-1 and para[i].text.find(groupmember)==-1:
            para[i].clear()
            try:
                j=1
                while para[i+j].text.find(groupmember+':')==-1:
                    para[i+j].clear()
                    j+=1
            except IndexError:
                continue
        elif para[i].text.find(groupmember)>-1:
            para[i].clear()
            
    null_list=[p.text for p in para]
    true_list = [i for i in null_list if i != '']
    data='\n'.join(true_list)
    D2.add_paragraph(when,style='Heading 1')
    D2.add_paragraph(data,style='Normal')
    D2.save(os.path.join(os.path.dirname(pathroad),savename+'.docx'))
            
    
    

    

