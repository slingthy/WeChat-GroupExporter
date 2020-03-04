#   Author : github.com/slingthy
#   Time : March 4, 2020
#   ---------------------------------------------
#   The purpose is to extract someone's record 
#   from wechat group and store it in a docx file.
#   !!!JUST 'TEXT' message.
#   ---------------------------------------------
#   Before RUN:
#   1. COPY the WeChat Group message
#   2. PASTE into Word document as 'xxxxxx.docx'
#   3. SAVE it, do NOT make any change
#   ---------------------------------------------
#   Dependencies : python-docx, lxml

from docx import Document
import time


when=time.strftime("%m.%d", time.localtime())
pathroad=input(r'Enter the path')+'\\'
name=input("Enter the filename( eg: xxx.docx )")
groupmember=input('Who you want?(only one member)')
savename=input("Save as?( eg: xxx.docx )")

D = Document(pathroad+name)
D2= Document()
para=D.paragraphs

#删除其他人的聊天记录
for i in range(0,len(para)):
    if para[i].text.find(":")>-1 and para[i].text.find(groupmember)==-1:
        para[i].clear()
        try:
            j=1
            while para[i+j].text.find(":")==-1:
                para[i+j].clear()
                j+=1
        except IndexError:
            continue
    elif para[i].text.find(groupmember)>-1:
        para[i].clear()
        
#存档于null_list
null_list=[p.text for p in para]

#清洗，去除空行
true_list = [i for i in null_list if i != '']
data='\n'.join(true_list)
D2.add_paragraph(when,style='Heading 1')
D2.add_paragraph(data,style='Normal')
print(when,data,sep='\n')
D2.save(pathroad+savename)
            
    
    

    

